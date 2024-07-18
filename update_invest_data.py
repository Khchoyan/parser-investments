import pandas as pd
import re

import docx
import aspose.words as aw
import os

import requests
from bs4 import BeautifulSoup as bs
import time
import datetime
from calendar import monthrange


def pars_year_by_months(year):
    '''
    Функция для получения ссылок на документы по месяцам.
    Для инвестиций реализовано возвращение названия последнего доступного месяца в конкретном году
    и ссылки на соответствующий раздел.
    '''
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }

    url = f'https://rosstat.gov.ru/storage/mediabank/Doklad_{year}.htm'
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links_1 = pd.DataFrame()
    for i in range(0, len(soup.find('table').find_all('tr')[1].find_all('tr')), 2):
        month_name = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[0].text
        month_name = month_name.replace('\n', '')
        if month_name.split()[-1].lower() == 'год':
            month_name = 'Январь-декабрь'
        dok_link = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[1].find_all('a')[0].get('href')
        if dok_link[:4] != 'http':
            dok_link = 'https://rosstat.gov.ru' + dok_link
        pril_link = soup.find('table').find_all('tr')[1].find_all('tr')[i + 1].find_all('td')[0].find_all('a')[0].get(
            'href')
        if pril_link[:4] != 'http':
            pril_link = 'https://rosstat.gov.ru' + pril_link
        links_1 = links_1._append([[month_name, dok_link, pril_link]])

    # Инвестиции выходят 4 раза в год
    links_1 = links_1[links_1[0].str.contains('февраль|апрель|июль|октябрь')]
    return links_1.iloc[0, 0].split('-')[-1], links_1.iloc[0, 1]


def str_month2digit_month(month):
    '''
    Функция переводит название месяца в его номер.
    '''
    month = month.strip().lower()
    if month == 'январь':
        return '01'
    elif month == 'февраль':
        return '02'
    elif month == 'март':
        return '03'
    elif month == 'апрель':
        return '04'
    elif month == 'май':
        return '05'
    elif month == 'июнь':
        return '06'
    elif month == 'июль':
        return '07'
    elif month == 'август':
        return '08'
    elif month == 'сентябрь':
        return '09'
    elif month == 'октябрь':
        return '10'
    elif month == 'ноябрь':
        return '11'
    elif month == 'декабрь':
        return '12'
    else:
        return 'unknown'


def download_document(year, month, url):
    '''
    Функция скачивает документ с данными по инвестициям за конкретный месяц.
    year - год в формате ХХХХ.
    month - полное название месяца на русском языке.
    url - ссылка на документ.
    Первые две переменные необходимы для назначения имени скачиваемому файлу.
    Возвращает путь к сохранённому файлу.
    '''
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links = pd.DataFrame()
    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        dok_link = link.get('href')
        links = links._append([[branch_name, dok_link]])

    indicator = 'Инвестиции в нефинансовые активы'
    if len(links[links[0] == indicator][1]) == 0:
        print(f'NO DOCUMENT {year}_{month}: {indicator}')
    else:
        link_to_download = links[links[0] == indicator][1].values[0]
        dok_name_to_download = f'{year}_{month}-2-4-0.doc'  # 2024_02-2-4-0.doc
        folder = os.getcwd()
        folder = os.path.join(folder, 'word_data', dok_name_to_download)

        response = requests.get(link_to_download, headers=header)
        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def doc2docx(path: str):
    '''
    Функция конвертирует документ формата .doc в формат .docx
    doc_path - абсолютный путь к документу
    '''
    doc = aw.Document(path)
    doc.save(path + 'x')
    print(f'Document {path} was converted to docx-format.')
    return path + 'x'


def reformat_date(date: str):
    '''
    Функция переформатирует даты
    '''
    date = date.strip()
    if date == 'I квартал':
        date = '31 march'
    elif date == 'I полугодие':
        date = '30 june'
    elif date == 'Январь-сентябрь':
        date = '30 september'
    elif date == 'Год':
        date = '31 december'
    return date


def parse_docx_document(path, year, month):
    '''
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    '''
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')

    data_table = [[] for _ in range(len(doc.tables[1].rows))]
    for i, row in enumerate(doc.tables[1].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    comment = data_table.iloc[-1, 0]

    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table = data_table[data_table.iloc[:, 0].str.contains(' I квартал|I полугодие|Январь-сентябрь|Год')]
    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    for i in range(len(data_table)):
        if ')' in data_table.iloc[i, 0]:
            data_table.iloc[i, 0] = data_table.iloc[i, 0][:-2]
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: reformat_date(x))

    if month == 'февраль':
        year -= 1
    for i in range(len(data_table)):
        if i < 4:
            data_table.iloc[i, 0] = pd.to_datetime(data_table.iloc[i, 0] + str(year - 1))
        else:
            data_table.iloc[i, 0] = pd.to_datetime(data_table.iloc[i, 0] + str(year))

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')
        try:
            data_table.iloc[:, i] = data_table.iloc[:, i].astype('float')
        except ValueError:
            print('parse_docx_document: Could not convert string to float. Unknown symbol.')
    print(f'Document {path} was parsed')

    return data_table, comment


def parse_all_docx_documents():
    '''
    Функция осуществляет парсинг всех документов в папке word_data.
    Данная функция сделана на всякий случай.
    '''
    paths = []
    folder = os.getcwd()
    folder = os.path.join(folder, 'word_data')
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('doc') and not file.startswith('~'):
                paths.append(os.path.join(root, file))
    paths = paths[:9]  # 9 - считывание файлов только до 2023 года включительно.
    # При увеличении числа считываемых файлов необходимо добавлять ещё годы в список ниже.
    all_data = pd.DataFrame()
    for doc_path_, year in zip(paths, [2015, 2016, 2017, 2018, 2019, 2020, 2021, 2022, 2023]):
        path = doc2docx(doc_path_)
        df, comm = parse_docx_document(path, year=year, month='апрель')
        os.remove(path)

        temp = df.iloc[:, :3]
        temp['3'] = comm
        temp.columns = ['Дата',
                        'Сумма инвестиций накопительным итогом, млрд рублей',
                        'Динамика инвестиций накопительным итогом, % к соответствующему периоду предыдущего года',
                        'Комментарий']
        all_data = all_data._append(temp)
    return all_data.drop_duplicates(subset=['Дата'])


def update_csv(data, csv_path='invest.csv'):
    """
    Функция осуществляет обновление файла с данными по инвестициям
    """
    data_csv = pd.read_csv(csv_path)
    data_csv = data_csv._append(data)
    data_csv = data_csv.drop_duplicates(subset=['Дата'], keep='last')
    data_csv.to_csv(csv_path, index=False)
    print(f'{csv_path} was apdated')


def create_new_date(last_date_in_file_year, last_date_in_file_month):
    now = datetime.datetime.now()
    lst_date = []
    _, last_day = monthrange(now.year, now.month)
    last_date = datetime.datetime.strptime(f"{now.year}-{now.month}-{last_day}", "%Y-%m-%d").date()

    for i in range((last_date.year - last_date_in_file_year) * 12 + last_date.month - last_date_in_file_month - 1):
        if last_date.month - 1 != 0:
            _, last_day = monthrange(last_date.year, last_date.month - 1)
            last_date = datetime.datetime.strptime(f"{last_date.year}-{last_date.month - 1}-{last_day}", "%Y-%m-%d").date()
        else:
            _, last_day = monthrange(last_date.year - 1, 12)
            last_date = datetime.datetime.strptime(f"{last_date.year - 1}-{12}-{last_day}", "%Y-%m-%d").date()
        lst_date.append(last_date)
    return sorted(lst_date)


def append_date_rez_file_Y(xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет дабавление месяцев, если их нет в файле.
    """
    data_xlsx = pd.read_excel(xlsx_path)
    year = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).year
    month = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).month
    date_lst = create_new_date(year, month)
    for date in date_lst:
        new_string = {'Целевой показатель': date}
        new_string.update({c: None for c in data_xlsx.columns[1:]})
        data_xlsx = data_xlsx._append(new_string, ignore_index=True)
    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_y(data, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if data.values[-1][0] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = data.columns[1]
    name_2 = data.columns[2]
    for j in data.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_1] = j[1]
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_2] = j[2]

    data_xlsx.to_excel(xlsx_path, index=False)


def main():
    '''
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    '''
    now = datetime.datetime.now().year
    # last_year_in_table = pd.to_datetime(pd.read_csv('invest.csv')['Дата'].iloc[-1]).year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).year

    if now - last_year_in_table < 2:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table + 1, now + 1):
            years.append(y)

    new_data = pd.DataFrame()
    for year in years:
        time.sleep(15)
        month, url = pars_year_by_months(year)
        print(month, url)
        time.sleep(15)
        path_to_docfile = download_document(year, month, url)

        path = doc2docx(path_to_docfile)
        df, comm = parse_docx_document(path, year=year, month=month)
        os.remove(path)

        temp = df.iloc[:, :3]
        temp['3'] = comm
        temp.columns = ['Дата',
                        'Сумма инвестиций накопительным итогом, млрд рублей',
                        'Динамика инвестиций накопительным итогом, % к соответствующему периоду предыдущего года',
                        'Комментарий']
        new_data = new_data._append(temp)

    new_data = new_data.drop_duplicates(subset=['Дата'], keep='last')
    update_csv(new_data, csv_path='invest.csv')

    del new_data['Комментарий']
    new_data = new_data.rename(columns={'Дата': 'Целевой показатель',
                                        'Сумма инвестиций накопительным итогом, млрд рублей':
                                            'Инвестиции в основной капитал накопленным итогом, млрд руб',
                                        'Динамика инвестиций накопительным итогом, % к соответствующему периоду '
                                        'предыдущего года':
                                            'Инвестиции, % накопленным итогом год к году'})
    update_rez_file_y(new_data, xlsx_path='rez_file_Y_v2.xlsx')


if __name__ == '__main__':
    main()
